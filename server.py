from fastapi import FastAPI, APIRouter, HTTPException, UploadFile, File, Form
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone
import re
import cloudinary
import cloudinary.uploader
from docx import Document
import io
import secrets

ROOT_DIR = Path(__file__).parent
from dotenv import load_dotenv
load_dotenv(ROOT_DIR / '.env')

# Configure Cloudinary FIRST
try:
    cloudinary.config(
        cloud_name=os.getenv("CLOUDINARY_CLOUD_NAME"),
        api_key=os.getenv("CLOUDINARY_API_KEY"),
        api_secret=os.getenv("CLOUDINARY_API_SECRET")
    )
    print("✓ Cloudinary configured successfully")
except Exception as e:
    print(f"✗ Cloudinary config error: {e}")

# MongoDB connection
mongo_url = os.environ.get('MONGO_URL')
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ.get('DB_NAME')]

# Create the main app
app = FastAPI()
api_router = APIRouter(prefix="/api")

# Models
class AdminLogin(BaseModel):
    username: str
    password: str

class ExamConfig(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    branch: str
    year: str
    semester: str
    subject: str
    num_students: int
    time_limit: int
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    questions_uploaded: bool = False
    questions_count: int = 0
    questions_per_student: int = 0
    sections: List[Dict[str, Any]] = []

class ExamConfigCreate(BaseModel):
    branch: str
    year: str
    semester: str
    subject: str
    num_students: int
    time_limit: int

class Question(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    exam_id: str
    question_number: int
    question_text: str
    has_code: bool = False
    code_snippet: Optional[str] = None
    options: List[Dict[str, str]]
    correct_answer: str
    section_id: Optional[str] = None
    image_url: Optional[str] = None  # Cloudinary URL
    image_public_id: Optional[str] = None  # Cloudinary public_id for deletion

class StudentRegister(BaseModel):
    name: str
    roll_number: str
    year: str
    semester: str
    branch: str
    section: str
    email: Optional[str] = None

class StudentLogin(BaseModel):
    roll_number: str
    password: str

class StudentLoginResponse(BaseModel):
    success: bool
    student: dict
    session_id: str
    requires_device_confirmation: bool = False
    existing_device_info: Optional[dict] = None

class DeviceConfirmation(BaseModel):
    roll_number: str
    password: str
    confirm_continue: bool

class ExamAttempt(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    student_id: str
    exam_id: str
    started_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    submitted_at: Optional[datetime] = None
    answers: Dict[str, str] = {}
    marked_for_review: List[str] = []
    completed: bool = False
    score: Optional[int] = None
    total_questions: Optional[int] = None
    suspicious_activity_count: int = 0

class AnswerSubmit(BaseModel):
    question_id: str
    answer: str

class SectionConfig(BaseModel):
    name: str
    question_ids: List[str]

# Admin Routes
@api_router.post("/admin/login")
async def admin_login(data: AdminLogin):
    if data.username == "admin" and data.password == "admin@4456":
        return {"success": True, "message": "Login successful"}
    raise HTTPException(status_code=401, detail="Invalid credentials")

@api_router.post("/admin/exam-config", response_model=ExamConfig)
async def create_exam_config(config: ExamConfigCreate):
    exam = ExamConfig(**config.model_dump())
    doc = exam.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    await db.exams.insert_one(doc)
    return exam

@api_router.get("/admin/exam-configs", response_model=List[ExamConfig])
async def get_exam_configs():
    exams = await db.exams.find({}, {"_id": 0}).to_list(1000)
    for exam in exams:
        if isinstance(exam.get('created_at'), str):
            exam['created_at'] = datetime.fromisoformat(exam['created_at'])
        # Ensure these fields exist
        if 'questions_per_student' not in exam:
            exam['questions_per_student'] = 0
        if 'questions_count' not in exam:
            exam['questions_count'] = 0
    return exams

@api_router.post("/admin/upload-questions/{exam_id}")
async def upload_questions(exam_id: str, file: UploadFile = File(...)):
    """Upload questions from DOCX with image extraction"""
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only DOCX files allowed")
    
    try:
        content = await file.read()
        doc = Document(io.BytesIO(content))
        
        questions = []
        question_images = {}  # Map question number to image URL
        
        # Extract all images first and map to question numbers
        try:
            image_counter = 0
            for para in doc.paragraphs:
                # Check if this paragraph has images
                for run in para.runs:
                    if hasattr(run._element, 'drawing_lst'):
                        for drawing in run._element.drawing_lst:
                            try:
                                blip = drawing.graphic.graphicData.pic.blipFill.blip
                                rId = blip.embed
                                image_part = doc.part.related_part(rId)
                                image_bytes = image_part.blob
                                
                                # Upload to Cloudinary
                                public_id = f"exam_questions/{exam_id}/img_{image_counter}"
                                upload_result = cloudinary.uploader.upload(
                                    image_bytes,
                                    public_id=public_id,
                                    overwrite=True,
                                    resource_type="auto",
                                    folder="exam_questions"
                                )
                                
                                # Try to get question number from paragraph text
                                para_text = para.text.strip()
                                q_match = re.match(r'^Q(\d+)', para_text, re.IGNORECASE)
                                if q_match:
                                    q_num = int(q_match.group(1))
                                    question_images[q_num] = upload_result['secure_url']
                                    print(f"✓ Image {image_counter} mapped to Q{q_num}: {upload_result['secure_url']}")
                                
                                image_counter += 1
                            except Exception as e:
                                print(f"✗ Error extracting image: {e}")
                                continue
        except Exception as e:
            print(f"Image extraction error: {e}")
        
        # Parse questions
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Match question pattern Q1, Q2, etc
            q_match = re.match(r'^Q(\d+)[:.]?\s*(.+)', text, re.DOTALL | re.IGNORECASE)
            if not q_match:
                continue
            
            question_number = int(q_match.group(1))
            full_text = q_match.group(2)
            
            lines = full_text.split('\n')
            question_lines = []
            options = []
            in_options = False
            code_lines = []
            in_code = False
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check for options (A), B), C), D))
                opt_match = re.match(r'^([A-D])\)\s*(.+)$', line, re.IGNORECASE)
                if opt_match:
                    in_options = True
                    in_code = False
                    letter = opt_match.group(1).upper()
                    value = opt_match.group(2).strip()
                    options.append({'letter': letter, 'value': value})
                    continue
                
                # Check for code
                if not in_options:
                    code_indicators = ['#include', 'int main', 'def ', 'public class', 'console.log', 
                                     'printf', 'System.out', 'cout', 'print(', 'using namespace', '```']
                    if any(indicator in line for indicator in code_indicators):
                        in_code = True
                    
                    if in_code:
                        code_lines.append(line)
                    else:
                        question_lines.append(line)
            
            if not options:
                continue
            
            question_text = ' '.join(question_lines)
            
            # Find correct answer
            correct_answer = None
            for opt in options:
                if '*' in opt['value']:
                    correct_answer = opt['letter']
                    opt['value'] = opt['value'].replace('*', '').strip()
            
            # Get image for this question number
            image_url = question_images.get(question_number)
            
            questions.append({
                'exam_id': exam_id,
                'question_number': question_number,
                'question_text': question_text,
                'has_code': bool(code_lines),
                'code_snippet': '\n'.join(code_lines) if code_lines else None,
                'options': options,
                'correct_answer': correct_answer,
                'id': str(uuid.uuid4()),
                'image_url': image_url,  # Cloudinary secure_url
                'created_at': datetime.now(timezone.utc).isoformat()
            })
        
        if not questions:
            raise HTTPException(status_code=400, detail="No valid questions found in document")
        
        # Insert questions
        await db.questions.insert_many(questions)
        
        # Update exam
        total_questions = len(questions)
        await db.exams.update_one(
            {"id": exam_id},
            {"$set": {
                "questions_uploaded": True,
                "questions_count": total_questions,
                "questions_per_student": total_questions
            }}
        )
        
        print(f"✓ Uploaded {len(questions)} questions with {len(question_images)} images")
        return {
            "success": True,
            "questions_count": len(questions),
            "images_count": len(question_images)
        }
    
    except Exception as e:
        print(f"✗ Upload error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error uploading questions: {str(e)}")

def extract_code(text):
    """Extract code snippets from question text"""
    # Look for code blocks with ``` markers
    code_match = re.search(r'```([^`]+)```', text, re.DOTALL)
    if code_match:
        return code_match.group(1).strip()
    
    # Fallback: Look for code patterns
    lines = text.split('\n')
    code_lines = []
    in_code = False
    
    for line in lines:
        code_indicators = ['#include', 'def ', 'public class', 'int main', 'console.log', 'printf', 'System.out', 'cout <<', 'print(']
        if any(indicator in line for indicator in code_indicators):
            in_code = True
        if in_code:
            code_lines.append(line)
            # End code block on closing brace or return statement
            if line.strip().endswith('}') or 'return 0;' in line or line.strip() == '':
                if line.strip().endswith('}'):
                    in_code = False
    
    return '\n'.join(code_lines).strip() if code_lines else None

@api_router.post("/admin/configure-question-count/{exam_id}")
async def configure_question_count(exam_id: str, count: int = Form(...)):
    await db.exams.update_one(
        {"id": exam_id},
        {"$set": {"questions_per_student": count}}
    )
    return {"success": True}

@api_router.post("/admin/organize-sections/{exam_id}")
async def organize_sections(exam_id: str, sections: List[dict]):
    """Save organized sections and update questions_per_student"""
    try:
        exam = await db.exams.find_one({"id": exam_id})
        if not exam:
            raise HTTPException(status_code=404, detail="Exam not found")
        
        # Count total questions assigned to sections
        total_assigned = 0
        for section in sections:
            total_assigned += len(section.get('question_ids', []))
        
        # Update exam with sections and update questions_per_student
        await db.exams.update_one(
            {"id": exam_id},
            {"$set": {
                "sections": sections,
                "questions_per_student": total_assigned  # Update to match selected questions
            }}
        )
        
        # Update question documents with section_id
        for section_idx, section in enumerate(sections):
            for question_id in section.get('question_ids', []):
                await db.questions.update_one(
                    {"id": question_id, "exam_id": exam_id},
                    {"$set": {"section_id": section_idx}}
                )
        
        # Mark unassigned questions (remove section_id)
        all_assigned_ids = []
        for section in sections:
            all_assigned_ids.extend(section.get('question_ids', []))
        
        await db.questions.update_many(
            {"exam_id": exam_id, "id": {"$nin": all_assigned_ids}},
            {"$unset": {"section_id": ""}}
        )
        
        return {"success": True, "questions_assigned": total_assigned}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error saving sections: {str(e)}")

@api_router.get("/admin/questions/{exam_id}")
async def get_exam_questions(exam_id: str):
    questions = await db.questions.find({"exam_id": exam_id}, {"_id": 0}).to_list(1000)
    return questions

@api_router.get("/admin/results")
async def get_results(
    branch: Optional[str] = None,
    year: Optional[str] = None,
    semester: Optional[str] = None,
    subject: Optional[str] = None,
    section: Optional[str] = None
):
    """Get all results - only count sectioned questions"""
    attempts = await db.exam_attempts.find({"completed": True}, {"_id": 0}).to_list(1000)
    
    results = []
    for attempt in attempts:
        student = await db.students.find_one({"id": attempt["student_id"]}, {"_id": 0})
        exam = await db.exams.find_one({"id": attempt["exam_id"]}, {"_id": 0})
        
        if not student or not exam:
            continue
        
        # Apply filters
        if branch and exam.get('branch') != branch:
            continue
        if year and exam.get('year') != year:
            continue
        if semester and exam.get('semester') != semester:
            continue
        if subject and exam.get('subject') != subject:
            continue
        if section and student.get('section') != section:
            continue
        
        # Get ONLY sectioned questions for this exam
        sectioned_questions = await db.questions.find({
            "exam_id": exam['id'],
            "section_id": {"$exists": True, "$ne": None}
        }).to_list(1000)
        
        total_sectioned = len(sectioned_questions);
        
        # Count only sectioned questions in score
        answers = attempt.get('answers', {})
        score = 0
        for q in sectioned_questions:
            if answers.get(q['id']) == q['correct_answer']:
                score += 1
        
        percentage = round((score / total_sectioned * 100) if total_sectioned > 0 else 0, 2)
        
        results.append({
            "id": attempt['id'],
            "roll_number": student['roll_number'],
            "student_name": student['name'],
            "subject": exam['subject'],
            "score": f"{score}/{total_sectioned}",
            "percentage": percentage,
            "date": attempt.get('submitted_at', attempt.get('started_at')),
            "branch": exam.get('branch'),
            "year": exam.get('year'),
            "semester": exam.get('semester'),
            "section": student.get('section')
        })
    
    return results

@api_router.get("/admin/results/{exam_id}")
async def get_exam_results(exam_id: str):
    """Get exam results with correct score calculation - only count sectioned questions"""
    try:
        exam = await db.exams.find_one({"id": exam_id}, {"_id": 0})
        if not exam:
            raise HTTPException(status_code=404, detail="Exam not found")
        
        attempts = await db.exam_attempts.find(
            {"exam_id": exam_id, "completed": True},
            {"_id": 0}
        ).to_list(1000)
        
        # Get all questions that are assigned to sections
        sectioned_questions = await db.questions.find({
            "exam_id": exam_id,
            "section_id": {"$exists": True, "$ne": None}
        }).to_list(1000)
        
        sectioned_question_ids = [q['id'] for q in sectioned_questions]
        total_sectioned = len(sectioned_question_ids);
        
        results = []
        for attempt in attempts:
            # Get student info
            student = await db.students.find_one(
                {"id": attempt["student_id"]},
                {"_id": 0, "password": 0}
            )
            
            if student:
                # Count only answers for sectioned questions
                answers = attempt.get('answers', {})
                score = 0
                
                # Get sectioned questions with correct answers
                for q in sectioned_questions:
                    if answers.get(q['id']) == q['correct_answer']:
                        score += 1
                
                percentage = round((score / total_sectioned * 100) if total_sectioned > 0 else 0, 2)
                
                results.append({
                    "student_name": student.get("name", "N/A"),
                    "roll_number": student.get("roll_number", "N/A"),
                    "score": score,
                    "total_questions": total_sectioned,
                    "percentage": percentage,
                    "time_taken": attempt.get("time_taken", 0),
                    "submitted_at": attempt.get("submitted_at", "N/A")
                })
        
        return {
            "exam": exam,
            "results": results,
            "total_students": len(results),
            "total_questions": total_sectioned
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching results: {str(e)}")

# Student Routes
@api_router.post("/student/register")
async def student_register(data: StudentRegister):
    # Check roll number length
    if len(data.roll_number) != 10:
        raise HTTPException(status_code=400, detail="Roll number must be 10 characters")
    
    # Check if roll number is alphanumeric
    if not data.roll_number.isalnum():
        raise HTTPException(status_code=400, detail="Roll number must be alphanumeric")
    
    # Check for duplicate roll number
    existing = await db.students.find_one({"roll_number": data.roll_number})
    if existing:
        raise HTTPException(status_code=400, detail="Roll number already exists")
    
    student = {
        "id": str(uuid.uuid4()),
        "name": data.name,
        "roll_number": data.roll_number,
        "year": data.year,
        "semester": data.semester,
        "branch": data.branch,
        "section": data.section,
        "email": data.email,
        "password": "Student@123",
        "active_session_id": None,
        "last_login_device": None,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.students.insert_one(student)
    return {"success": True, "message": "Registration successful"}

@api_router.post("/student/login")
async def student_login(data: StudentLogin):
    student = await db.students.find_one({"roll_number": data.roll_number}, {"_id": 0})
    if not student:
        raise HTTPException(status_code=401, detail="Invalid roll number or password")
    
    if student['password'] != data.password:
        raise HTTPException(status_code=401, detail="Invalid roll number or password")
    
    # Check if student has an active session
    active_session_id = student.get('active_session_id')
    
    if active_session_id:
        # Active session exists - return without creating new session
        return {
            "success": False,
            "requires_device_confirmation": True,
            "existing_device_info": {
                "device": student.get('last_login_device', 'Unknown Device'),
                "message": "You are already logged in on another device."
            },
            "student": {
                "id": student['id'],
                "name": student['name'],
                "roll_number": student['roll_number']
            }
        }
    
    # No active session - create new one
    new_session_id = secrets.token_urlsafe(32)
    
    await db.students.update_one(
        {"id": student['id']},
        {"$set": {
            "active_session_id": new_session_id,
            "last_login_device": "Device",
            "last_login_time": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    # Remove password from response for security
    student_response = {k: v for k, v in student.items() if k != 'password'}
    
    return {
        "success": True,
        "student": student_response,
        "session_id": new_session_id,
        "requires_device_confirmation": False
    }

@api_router.post("/student/confirm-device-login")
async def confirm_device_login(data: DeviceConfirmation):
    """Handle the device confirmation for duplicate login"""
    student = await db.students.find_one({"roll_number": data.roll_number}, {"_id": 0})
    if not student:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if student['password'] != data.password:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if not data.confirm_continue:
        # User clicked Cancel - return to login
        return {
            "success": False,
            "message": "Login cancelled. Please try again.",
            "cancelled": True
        }
    
    # User clicked Continue - invalidate old session and create new one
    new_session_id = secrets.token_urlsafe(32)
    
    await db.students.update_one(
        {"id": student['id']},
        {"$set": {
            "active_session_id": new_session_id,
            "last_login_device": "Device",
            "last_login_time": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    # Remove password from response
    student_response = {k: v for k, v in student.items() if k != 'password'}
    
    return {
        "success": True,
        "student": student_response,
        "session_id": new_session_id,
        "message": "Previous session has been logged out. You are now logged in."
    }

@api_router.post("/student/validate-session")
async def validate_session(session_id: str):
    """Validate if the session is still active"""
    # Find student with this session
    student = await db.students.find_one({"active_session_id": session_id}, {"_id": 0})
    
    if not student:
        raise HTTPException(status_code=401, detail="Session invalid or expired")
    
    return {"success": True, "student_id": student['id']}

@api_router.post("/student/logout")
async def student_logout(student_id: str):
    """Logout student and clear session"""
    result = await db.students.update_one(
        {"id": student_id},
        {"$set": {"active_session_id": None}}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Student not found")
    
    return {"success": True, "message": "Logged out successfully"}

# Admin Routes
@api_router.post("/admin/login")
async def admin_login(data: AdminLogin):
    if data.username == "admin" and data.password == "admin@4456":
        return {"success": True, "message": "Login successful"}
    raise HTTPException(status_code=401, detail="Invalid credentials")

@api_router.post("/admin/exam-config", response_model=ExamConfig)
async def create_exam_config(config: ExamConfigCreate):
    exam = ExamConfig(**config.model_dump())
    doc = exam.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    await db.exams.insert_one(doc)
    return exam

@api_router.get("/admin/exam-configs", response_model=List[ExamConfig])
async def get_exam_configs():
    exams = await db.exams.find({}, {"_id": 0}).to_list(1000)
    for exam in exams:
        if isinstance(exam.get('created_at'), str):
            exam['created_at'] = datetime.fromisoformat(exam['created_at'])
        # Ensure these fields exist
        if 'questions_per_student' not in exam:
            exam['questions_per_student'] = 0
        if 'questions_count' not in exam:
            exam['questions_count'] = 0
    return exams

@api_router.post("/admin/upload-questions/{exam_id}")
async def upload_questions(exam_id: str, file: UploadFile = File(...)):
    """Upload questions from DOCX with image extraction"""
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only DOCX files allowed")
    
    try:
        content = await file.read()
        doc = Document(io.BytesIO(content))
        
        questions = []
        question_images = {}  # Map question number to image URL
        
        # Extract all images first and map to question numbers
        try:
            image_counter = 0
            for para in doc.paragraphs:
                # Check if this paragraph has images
                for run in para.runs:
                    if hasattr(run._element, 'drawing_lst'):
                        for drawing in run._element.drawing_lst:
                            try:
                                blip = drawing.graphic.graphicData.pic.blipFill.blip
                                rId = blip.embed
                                image_part = doc.part.related_part(rId)
                                image_bytes = image_part.blob
                                
                                # Upload to Cloudinary
                                public_id = f"exam_questions/{exam_id}/img_{image_counter}"
                                upload_result = cloudinary.uploader.upload(
                                    image_bytes,
                                    public_id=public_id,
                                    overwrite=True,
                                    resource_type="auto",
                                    folder="exam_questions"
                                )
                                
                                # Try to get question number from paragraph text
                                para_text = para.text.strip()
                                q_match = re.match(r'^Q(\d+)', para_text, re.IGNORECASE)
                                if q_match:
                                    q_num = int(q_match.group(1))
                                    question_images[q_num] = upload_result['secure_url']
                                    print(f"✓ Image {image_counter} mapped to Q{q_num}: {upload_result['secure_url']}")
                                
                                image_counter += 1
                            except Exception as e:
                                print(f"✗ Error extracting image: {e}")
                                continue
        except Exception as e:
            print(f"Image extraction error: {e}")
        
        # Parse questions
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Match question pattern Q1, Q2, etc
            q_match = re.match(r'^Q(\d+)[:.]?\s*(.+)', text, re.DOTALL | re.IGNORECASE)
            if not q_match:
                continue
            
            question_number = int(q_match.group(1))
            full_text = q_match.group(2)
            
            lines = full_text.split('\n')
            question_lines = []
            options = []
            in_options = False
            code_lines = []
            in_code = False
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check for options (A), B), C), D))
                opt_match = re.match(r'^([A-D])\)\s*(.+)$', line, re.IGNORECASE)
                if opt_match:
                    in_options = True
                    in_code = False
                    letter = opt_match.group(1).upper()
                    value = opt_match.group(2).strip()
                    options.append({'letter': letter, 'value': value})
                    continue
                
                # Check for code
                if not in_options:
                    code_indicators = ['#include', 'int main', 'def ', 'public class', 'console.log', 
                                     'printf', 'System.out', 'cout', 'print(', 'using namespace', '```']
                    if any(indicator in line for indicator in code_indicators):
                        in_code = True
                    
                    if in_code:
                        code_lines.append(line)
                    else:
                        question_lines.append(line)
            
            if not options:
                continue
            
            question_text = ' '.join(question_lines)
            
            # Find correct answer
            correct_answer = None
            for opt in options:
                if '*' in opt['value']:
                    correct_answer = opt['letter']
                    opt['value'] = opt['value'].replace('*', '').strip()
            
            # Get image for this question number
            image_url = question_images.get(question_number)
            
            questions.append({
                'exam_id': exam_id,
                'question_number': question_number,
                'question_text': question_text,
                'has_code': bool(code_lines),
                'code_snippet': '\n'.join(code_lines) if code_lines else None,
                'options': options,
                'correct_answer': correct_answer,
                'id': str(uuid.uuid4()),
                'image_url': image_url,  # Cloudinary secure_url
                'created_at': datetime.now(timezone.utc).isoformat()
            })
        
        if not questions:
            raise HTTPException(status_code=400, detail="No valid questions found in document")
        
        # Insert questions
        await db.questions.insert_many(questions)
        
        # Update exam
        total_questions = len(questions)
        await db.exams.update_one(
            {"id": exam_id},
            {"$set": {
                "questions_uploaded": True,
                "questions_count": total_questions,
                "questions_per_student": total_questions
            }}
        )
        
        print(f"✓ Uploaded {len(questions)} questions with {len(question_images)} images")
        return {
            "success": True,
            "questions_count": len(questions),
            "images_count": len(question_images)
        }
    
    except Exception as e:
        print(f"✗ Upload error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error uploading questions: {str(e)}")

def extract_code(text):
    """Extract code snippets from question text"""
    # Look for code blocks with ``` markers
    code_match = re.search(r'```([^`]+)```', text, re.DOTALL)
    if code_match:
        return code_match.group(1).strip()
    
    # Fallback: Look for code patterns
    lines = text.split('\n')
    code_lines = []
    in_code = False
    
    for line in lines:
        code_indicators = ['#include', 'def ', 'public class', 'int main', 'console.log', 'printf', 'System.out', 'cout <<', 'print(']
        if any(indicator in line for indicator in code_indicators):
            in_code = True
        if in_code:
            code_lines.append(line)
            # End code block on closing brace or return statement
            if line.strip().endswith('}') or 'return 0;' in line or line.strip() == '':
                if line.strip().endswith('}'):
                    in_code = False
    
    return '\n'.join(code_lines).strip() if code_lines else None

@api_router.post("/admin/configure-question-count/{exam_id}")
async def configure_question_count(exam_id: str, count: int = Form(...)):
    await db.exams.update_one(
        {"id": exam_id},
        {"$set": {"questions_per_student": count}}
    )
    return {"success": True}

@api_router.post("/admin/organize-sections/{exam_id}")
async def organize_sections(exam_id: str, sections: List[dict]):
    """Save organized sections and update questions_per_student"""
    try:
        exam = await db.exams.find_one({"id": exam_id})
        if not exam:
            raise HTTPException(status_code=404, detail="Exam not found")
        
        # Count total questions assigned to sections
        total_assigned = 0
        for section in sections:
            total_assigned += len(section.get('question_ids', []))
        
        # Update exam with sections and update questions_per_student
        await db.exams.update_one(
            {"id": exam_id},
            {"$set": {
                "sections": sections,
                "questions_per_student": total_assigned  # Update to match selected questions
            }}
        )
        
        # Update question documents with section_id
        for section_idx, section in enumerate(sections):
            for question_id in section.get('question_ids', []):
                await db.questions.update_one(
                    {"id": question_id, "exam_id": exam_id},
                    {"$set": {"section_id": section_idx}}
                )
        
        # Mark unassigned questions (remove section_id)
        all_assigned_ids = []
        for section in sections:
            all_assigned_ids.extend(section.get('question_ids', []))
        
        await db.questions.update_many(
            {"exam_id": exam_id, "id": {"$nin": all_assigned_ids}},
            {"$unset": {"section_id": ""}}
        )
        
        return {"success": True, "questions_assigned": total_assigned}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error saving sections: {str(e)}")

@api_router.get("/admin/questions/{exam_id}")
async def get_exam_questions(exam_id: str):
    questions = await db.questions.find({"exam_id": exam_id}, {"_id": 0}).to_list(1000)
    return questions

@api_router.get("/admin/results")
async def get_results(
    branch: Optional[str] = None,
    year: Optional[str] = None,
    semester: Optional[str] = None,
    subject: Optional[str] = None,
    section: Optional[str] = None
):
    """Get all results - only count sectioned questions"""
    attempts = await db.exam_attempts.find({"completed": True}, {"_id": 0}).to_list(1000)
    
    results = []
    for attempt in attempts:
        student = await db.students.find_one({"id": attempt["student_id"]}, {"_id": 0})
        exam = await db.exams.find_one({"id": attempt["exam_id"]}, {"_id": 0})
        
        if not student or not exam:
            continue
        
        # Apply filters
        if branch and exam.get('branch') != branch:
            continue
        if year and exam.get('year') != year:
            continue
        if semester and exam.get('semester') != semester:
            continue
        if subject and exam.get('subject') != subject:
            continue
        if section and student.get('section') != section:
            continue
        
        # Get ONLY sectioned questions for this exam
        sectioned_questions = await db.questions.find({
            "exam_id": exam['id'],
            "section_id": {"$exists": True, "$ne": None}
        }).to_list(1000)
        
        total_sectioned = len(sectioned_questions);
        
        # Count only sectioned questions in score
        answers = attempt.get('answers', {})
        score = 0
        for q in sectioned_questions:
            if answers.get(q['id']) == q['correct_answer']:
                score += 1
        
        percentage = round((score / total_sectioned * 100) if total_sectioned > 0 else 0, 2)
        
        results.append({
            "id": attempt['id'],
            "roll_number": student['roll_number'],
            "student_name": student['name'],
            "subject": exam['subject'],
            "score": f"{score}/{total_sectioned}",
            "percentage": percentage,
            "date": attempt.get('submitted_at', attempt.get('started_at')),
            "branch": exam.get('branch'),
            "year": exam.get('year'),
            "semester": exam.get('semester'),
            "section": student.get('section')
        })
    
    return results

@api_router.get("/admin/results/{exam_id}")
async def get_exam_results(exam_id: str):
    """Get exam results with correct score calculation - only count sectioned questions"""
    try:
        exam = await db.exams.find_one({"id": exam_id}, {"_id": 0})
        if not exam:
            raise HTTPException(status_code=404, detail="Exam not found")
        
        attempts = await db.exam_attempts.find(
            {"exam_id": exam_id, "completed": True},
            {"_id": 0}
        ).to_list(1000)
        
        # Get all questions that are assigned to sections
        sectioned_questions = await db.questions.find({
            "exam_id": exam_id,
            "section_id": {"$exists": True, "$ne": None}
        }).to_list(1000)
        
        sectioned_question_ids = [q['id'] for q in sectioned_questions]
        total_sectioned = len(sectioned_question_ids);
        
        results = []
        for attempt in attempts:
            # Get student info
            student = await db.students.find_one(
                {"id": attempt["student_id"]},
                {"_id": 0, "password": 0}
            )
            
            if student:
                # Count only answers for sectioned questions
                answers = attempt.get('answers', {})
                score = 0
                
                # Get sectioned questions with correct answers
                for q in sectioned_questions:
                    if answers.get(q['id']) == q['correct_answer']:
                        score += 1
                
                percentage = round((score / total_sectioned * 100) if total_sectioned > 0 else 0, 2)
                
                results.append({
                    "student_name": student.get("name", "N/A"),
                    "roll_number": student.get("roll_number", "N/A"),
                    "score": score,
                    "total_questions": total_sectioned,
                    "percentage": percentage,
                    "time_taken": attempt.get("time_taken", 0),
                    "submitted_at": attempt.get("submitted_at", "N/A")
                })
        
        return {
            "exam": exam,
            "results": results,
            "total_students": len(results),
            "total_questions": total_sectioned
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching results: {str(e)}")

@api_router.post("/student/save-answer/{attempt_id}")
async def save_answer(attempt_id: str, data: AnswerSubmit):
    await db.exam_attempts.update_one(
        {"id": attempt_id},
        {"$set": {f"answers.{data.question_id}": data.answer}}
    )
    return {"success": True}

@api_router.post("/student/mark-review/{attempt_id}/{question_id}")
async def mark_for_review(attempt_id: str, question_id: str, mark: bool):
    attempt = await db.exam_attempts.find_one({"id": attempt_id})
    marked = attempt.get('marked_for_review', [])
    
    if mark and question_id not in marked:
        marked.append(question_id)
    elif not mark and question_id in marked:
        marked.remove(question_id)
    
    await db.exam_attempts.update_one(
        {"id": attempt_id},
        {"$set": {"marked_for_review": marked}}
    )
    return {"success": True}

@api_router.post("/student/suspicious-activity/{attempt_id}")
async def report_suspicious_activity(attempt_id: str):
    await db.exam_attempts.update_one(
        {"id": attempt_id},
        {"$inc": {"suspicious_activity_count": 1}}
    )
    return {"success": True}

@api_router.post("/student/submit-exam/{attempt_id}")
async def submit_exam(attempt_id: str):
    attempt = await db.exam_attempts.find_one({"id": attempt_id})
    if not attempt:
        raise HTTPException(status_code=404, detail="Attempt not found")
    
    # Get the SELECTED questions from attempt (not all exam questions)
    if attempt.get('question_ids'):
        questions = await db.questions.find({
            "id": {"$in": attempt['question_ids']}
        }, {"_id": 0}).to_list(1000)
    else:
        # Fallback to all exam questions if question_ids not saved
        questions = await db.questions.find(
            {"exam_id": attempt['exam_id']}, 
            {"_id": 0}
        ).to_list(1000)
    
    # Calculate score
    score = 0
    answers = attempt.get('answers', {})
    for question in questions:
        if answers.get(question['id']) == question['correct_answer']:
            score += 1
    
    # Calculate time taken
    started_at = datetime.fromisoformat(attempt['started_at']) if isinstance(attempt['started_at'], str) else attempt['started_at']
    time_taken = int((datetime.now(timezone.utc) - started_at).total_seconds())
    
    # Update attempt with correct total_questions
    await db.exam_attempts.update_one(
        {"id": attempt_id},
        {"$set": {
            "completed": True,
            "submitted_at": datetime.now(timezone.utc).isoformat(),
            "score": score,
            "total_questions": len(questions),  # Use selected questions count
            "time_taken": time_taken
        }}
    )
    
    return {"success": True, "score": score, "total": len(questions)}

@api_router.get("/student/results/{student_id}")
async def get_student_results(student_id: str):
    """Get student results - only count sectioned questions"""
    attempts = await db.exam_attempts.find({
        "student_id": student_id,
        "completed": True
    }, {"_id": 0}).to_list(1000)
    
    results = []
    for attempt in attempts:
        exam = await db.exams.find_one({"id": attempt['exam_id']}, {"_id": 0})
        if exam:
            # Get ONLY sectioned questions
            sectioned_questions = await db.questions.find({
                "exam_id": attempt['exam_id'],
                "section_id": {"$exists": True, "$ne": None}
            }).to_list(1000)
            
            total_sectioned = len(sectioned_questions);
            
            # Count only sectioned questions in score
            answers = attempt.get('answers', {})
            score = 0
            for q in sectioned_questions:
                if answers.get(q['id']) == q['correct_answer']:
                    score += 1
            
            percentage = round((score / total_sectioned * 100) if total_sectioned > 0 else 0, 2)
            
            results.append({
                "subject": exam['subject'],
                "branch": exam['branch'],
                "year": exam['year'],
                "semester": exam['semester'],
                "score": score,
                "total": total_sectioned,
                "percentage": percentage,
                "date": attempt.get('submitted_at')
            })
    
    return results

@api_router.get("/student/available-exams/{student_id}")
async def get_available_exams(student_id: str):
    """Get all available exams for a student based on their profile"""
    try:
        # Get student details
        student = await db.students.find_one({"id": student_id}, {"_id": 0})
        if not student:
            raise HTTPException(status_code=404, detail="Student not found")
        
        # Get all exams that match student's branch, year, and semester
        exams = await db.exams.find({
            "branch": student['branch'],
            "year": student['year'],
            "semester": student['semester']
        }, {"_id": 0}).to_list(1000)
        
        # For each exam, check if student has already completed it
        for exam in exams:
            completed_attempt = await db.exam_attempts.find_one({
                "student_id": student_id,
                "exam_id": exam['id'],
                "completed": True
            })
            exam['is_completed'] = completed_attempt is not None
        
        return exams
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching exams: {str(e)}")

@api_router.post("/student/start-exam/{exam_id}/{student_id}")
async def start_exam(exam_id: str, student_id: str, data: dict = None):
    """Start an exam and create an exam attempt for the student"""
    try:
        # Validate student exists
        student = await db.students.find_one({"id": student_id}, {"_id": 0})
        if not student:
            raise HTTPException(status_code=404, detail="Student not found")
        
        # Validate exam exists
        exam = await db.exams.find_one({"id": exam_id}, {"_id": 0})
        if not exam:
            raise HTTPException(status_code=404, detail="Exam not found")
        
        # Check if student already has a completed attempt
        completed_attempt = await db.exam_attempts.find_one({
            "student_id": student_id,
            "exam_id": exam_id,
            "completed": True
        })
        if completed_attempt:
            raise HTTPException(status_code=400, detail="You have already completed this exam")
        
        # Get sectioned questions for the exam
        sectioned_questions = await db.questions.find({
            "exam_id": exam_id,
            "section_id": {"$exists": True, "$ne": None}
        }, {"_id": 0}).to_list(1000)
        
        if not sectioned_questions:
            raise HTTPException(status_code=400, detail="No questions available for this exam")
        
        # Get question IDs only
        question_ids = [q['id'] for q in sectioned_questions]
        
        # Create new exam attempt
        attempt = ExamAttempt(
            student_id=student_id,
            exam_id=exam_id
        )
        
        attempt_doc = attempt.model_dump()
        attempt_doc['started_at'] = attempt_doc['started_at'].isoformat()
        attempt_doc['question_ids'] = question_ids  # Store the questions for this attempt
        
        await db.exam_attempts.insert_one(attempt_doc)
        
        return {
            "success": True,
            "attempt_id": attempt.id,
            "exam": exam,
            "questions": sectioned_questions,
            "time_limit": exam.get('time_limit', 0)
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error starting exam: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error starting exam: {str(e)}")

@api_router.delete("/admin/exam-config/{exam_id}")
async def delete_exam_config(exam_id: str):
    """Delete exam and all associated questions"""
    try:
        exam = await db.exams.find_one({"id": exam_id})
        if not exam:
            raise HTTPException(status_code=404, detail="Exam not found")
        
        # delete Cloudinary images associated with exam questions
        questions_with_images = await db.questions.find({"exam_id": exam_id, "image_public_id": {"$exists": True, "$ne": None}}).to_list(1000)
        for q in questions_with_images:
            public_id = q.get("image_public_id")
            if public_id:
                try:
                    cloudinary.uploader.destroy(public_id, resource_type='image')
                except Exception as e:
                    # don't fail whole request on cloudinary error; log for diagnostics
                    print(f"Warning: failed to destroy Cloudinary image {public_id}: {e}")
        
        # delete all questions and attempts
        await db.questions.delete_many({"exam_id": exam_id})
        await db.exam_attempts.delete_many({"exam_id": exam_id})
        result = await db.exams.delete_one({"id": exam_id})
        
        if result.deleted_count == 0:
            raise HTTPException(status_code=500, detail="Failed to delete exam")
        
        return {"success": True, "message": "Exam deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting exam: {str(e)}")

@api_router.post("/admin/upload-question-image/{exam_id}/{question_id}")
async def upload_question_image(exam_id: str, question_id: str, file: UploadFile = File(...)):
    """Upload image for a specific question to Cloudinary"""
    # Validate file type
    allowed_types = ['image/png', 'image/jpeg', 'image/jpg', 'image/gif', 'image/bmp']
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Only PNG, JPG, GIF, BMP images allowed")
    
    try:
        content = await file.read()
        
        # Generate public_id for better organization
        public_id = f"exam_questions/{exam_id}/{question_id}"
        
        # Upload to Cloudinary
        upload_result = cloudinary.uploader.upload(
            content,
            public_id=public_id,
            overwrite=True,
            resource_type="auto",
            folder="exam_questions"
        )
        
        image_url = upload_result['secure_url']
        
        # Update question with image URL
        result = await db.questions.update_one(
            {"id": question_id, "exam_id": exam_id},
            {"$set": {
                "image_url": image_url,
                "image_public_id": upload_result['public_id'],
                "updated_at": datetime.now(timezone.utc).isoformat()
            }}
        )
        
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Question not found")
        
        return {
            "success": True, 
            "image_url": image_url,
            "message": "Image uploaded successfully to Cloudinary"
        }
    
    except Exception as e:
        print(f"Error uploading to Cloudinary: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error uploading image: {str(e)}")

@api_router.delete("/admin/remove-question-image/{exam_id}/{question_id}")
async def remove_question_image(exam_id: str, question_id: str):
    """Remove image from a question"""
    try:
        # Get question to find public_id
        question = await db.questions.find_one({"id": question_id, "exam_id": exam_id})
        
        if not question:
            raise HTTPException(status_code=404, detail="Question not found")
        
        # Delete from Cloudinary if public_id exists
        if question.get('image_public_id'):
            try:
                cloudinary.uploader.destroy(question['image_public_id'])
            except Exception as e:
                print(f"Error deleting from Cloudinary: {e}")
        
        # Remove image from question in database
        await db.questions.update_one(
            {"id": question_id, "exam_id": exam_id},
            {"$set": {
                "image_url": None,
                "image_public_id": None
            }}
        )
        
        return {"success": True, "message": "Image removed successfully"}
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error removing image: {str(e)}")

# Add endpoint to get image from database if needed
@api_router.get("/admin/question-image/{exam_id}/{question_id}")
async def get_question_image(exam_id: str, question_id: str):
    """Get image from database"""
    question = await db.questions.find_one(
        {"id": question_id, "exam_id": exam_id},
        {"image_base64": 1, "_id": 0}
    )
    
    if not question or not question.get('image_base64'):
        raise HTTPException(status_code=404, detail="Image not found")
    
    return {
        "image_base64": question['image_base64']
    }

@api_router.get("/admin/sections/{exam_id}")
async def get_sections(exam_id: str):
    """Get all sections for an exam"""
    try:
        exam = await db.exams.find_one({"id": exam_id}, {"_id": 0})
        if not exam:
            raise HTTPException(status_code=404, detail="Exam not found")
        
        # Sections are stored in exam.sections array
        sections = exam.get('sections', [])
        print(f"✓ Returning {len(sections)} sections for exam {exam_id}")
        return sections
    except Exception as e:
        print(f"✗ Error fetching sections: {e}")
        raise HTTPException(status_code=500, detail=f"Error fetching sections: {str(e)}")

@api_router.get("/admin/exam-status/{exam_id}")
async def get_exam_status(exam_id: str):
    """Get exam status including questions count and sections"""
    try:
        exam = await db.exams.find_one({"id": exam_id}, {"_id": 0})
        if not exam:
            raise HTTPException(status_code=404, detail="Exam not found")
        
        # Count all questions
        all_questions = await db.questions.count_documents({"exam_id": exam_id})
        
        # Count sectioned questions
        sectioned_questions = await db.questions.count_documents({
            "exam_id": exam_id,
            "section_id": {"$exists": True, "$ne": None}
        })
        
        # Get sections
        sections = exam.get('sections', [])
        
        return {
            "total_questions": all_questions,
            "sectioned_questions": sectioned_questions,
            "unassigned_questions": all_questions - sectioned_questions,
            "sections_count": len(sections),
            "sections": sections,
            "questions_per_student": exam.get('questions_per_student', all_questions)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching exam status: {str(e)}")

@api_router.get("/student/exam/{exam_id}/{attempt_id}")
async def get_exam_data(exam_id: str, attempt_id: str):
    """Fetch exam data and questions for a specific attempt"""
    try:
        # Get exam config
        exam = await db.exams.find_one({"id": exam_id})
        if not exam:
            raise HTTPException(status_code=404, detail="Exam not found")
        
        # Get exam attempt
        attempt = await db.exam_attempts.find_one({"id": attempt_id})
        if not attempt:
            raise HTTPException(status_code=404, detail="Attempt not found")
        
        # Get questions for this attempt
        questions = await db.questions.find({
            "exam_id": exam_id,
            "$or": [
                {"section_id": {"$in": [q for q in attempt.get("questions", [])]}},
                {"id": {"$in": attempt.get("questions", [])}}
            ]
        }).to_list(length=None)
        
        # If no questions found by section, get all exam questions
        if not questions:
            questions = await db.questions.find({"exam_id": exam_id}).to_list(length=None)
        
        # Organize questions by sections if they exist
        sections = exam.get("sections", [])
        
        return {
            "exam": {
                "id": exam["id"],
                "subject": exam.get("subject"),
                "branch": exam.get("branch"),
                "year": exam.get("year"),
                "semester": exam.get("semester"),
                "time_limit": exam.get("time_limit"),
                "sections": sections,
                "total_questions": len(questions)
            },
            "questions": [
                {
                    "id": q["id"],
                    "question_number": q.get("question_number"),
                    "question_text": q.get("question_text"),
                    "options": q.get("options", []),
                    "has_code": q.get("has_code", False),
                    "code_snippet": q.get("code_snippet"),
                    "image_url": q.get("image_url"),
                    "section_id": q.get("section_id")
                }
                for q in questions
            ],
            "attempt": {
                "id": attempt["id"],
                "answers": attempt.get("answers", {}),
                "marked_for_review": attempt.get("marked_for_review", [])
            }
        }
    except Exception as e:
        logger.error(f"Error fetching exam data: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# ...existing code...

# IMPORTANT: Include the router in the app
app.include_router(api_router)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"]
)

@app.get("/")
async def root():
    return {"message": "Exam Browser API"}

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()

# Add this at the very end to run the server
if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)